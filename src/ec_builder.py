"""
ec_builder.py
Fills MASTER_TEMPLATE.docx with claim data to produce the Essentiality
Certificate.

NOTE on formatting fidelity: the master template's header block is one
merged table cell containing several field labels in a single paragraph
(designed for manual typewriter-style filling). To keep programmatic
filling reliable, this builder rewrites that block as clearly labelled
lines ("Name of Claimant: ...", "Period of Treatment: ...", etc.) using
the SAME font/size/bold as the original paragraph, rather than trying to
splice values into the exact original tab positions. Everything else
(item table, section order, ditto rule, signature blocks, fixed
declaration page) is preserved from the template unchanged.
"""

from __future__ import annotations
import copy
import io
from datetime import datetime
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

SECTION_ORDER = ["final_bill", "hc_medicine", "medicine", "discharge_medicine", "test", "misc"]
SECTION_LABELS = {
    "hc_medicine": "MEDICINE BILL(H.C.)",
    "medicine": "MEDICINE BILL",
    "discharge_medicine": "DISCH. MEDICINE",
}


def _fmt_date(d: Optional[str]) -> str:
    return d or "……………"


def _sort_key(item):
    try:
        return datetime.strptime(item.get("date") or "01-01-1900", "%d-%m-%Y")
    except ValueError:
        return datetime(1900, 1, 1)


def _build_ec_rows(items: list[dict]) -> list[dict]:
    """Apply EC section order (final_bill -> hc_medicine -> medicine ->
    discharge_medicine -> test -> misc), exclude return_credit, sort each
    section by date ascending, and compute the display name per row
    (regular medicines use the raw item name, category-labelled sections
    use the fixed label)."""
    usable = [i for i in items if i.get("category") != "return_credit"]
    rows = []
    for cat in SECTION_ORDER:
        section_items = sorted([i for i in usable if i["category"] == cat], key=_sort_key)
        for it in section_items:
            display_name = it["name"]
            rows.append({
                "name": display_name,
                "bill_no": it.get("bill_no") or "",
                "date": it.get("date") or "",
                "amount": it.get("gross_amount"),
            })
    return rows


def _apply_ditto(rows: list[dict]) -> list[dict]:
    prev_bill, prev_date = None, None
    for r in rows:
        bill, date = r["bill_no"], r["date"]
        r["bill_no_display"] = ",," if bill and bill == prev_bill else bill
        r["date_display"] = ",," if date and date == prev_date else date
        prev_bill, prev_date = bill, date
    return rows


def _set_cell_text(cell, text: str, base_run=None):
    """Replace a table cell's text while keeping (roughly) the formatting
    of its first existing run, or of base_run if the cell is empty."""
    para = cell.paragraphs[0]
    fmt_run = para.runs[0] if para.runs else base_run
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for extra_p in cell.paragraphs[1:]:
        extra_p._element.getparent().remove(extra_p._element)
    run = para.add_run(str(text) if text is not None else "")
    if fmt_run is not None:
        run.font.name = fmt_run.font.name
        run.font.size = fmt_run.font.size
        run.font.bold = fmt_run.font.bold
    return run


def _clone_row(table, template_row_index=0):
    tbl = table._tbl
    src_tr = table.rows[template_row_index]._tr
    new_tr = copy.deepcopy(src_tr)
    tbl.append(new_tr)
    return table.rows[-1]


def _find_item_table(doc):
    """Locate the table whose header row contains 'Sr. No.' and
    'Name & Quantity' - this is the item table to populate."""
    for t in doc.tables:
        header_text = " ".join(c.text for c in t.rows[0].cells).upper()
        if "SR. NO" in header_text and "NAME" in header_text:
            return t
    raise RuntimeError("Could not locate item table in MASTER_TEMPLATE.docx")


def _find_header_table(doc):
    for t in doc.tables:
        if "NAME OF CLAIMANT" in t.rows[0].cells[0].text.upper():
            return t
    raise RuntimeError("Could not locate header/info table in MASTER_TEMPLATE.docx")


def build_ec_docx(template_path: str, claim: dict, claimant_info: dict) -> io.BytesIO:
    doc = Document(template_path)

    # ---------- Header block ----------
    header_table = _find_header_table(doc)
    admission = claim.get("admission", {})
    patient = claim.get("patient", {})
    hospital = claim.get("hospital", {})

    header_lines = [
        f"Name of Claimant: {claimant_info.get('claimant_name') or ''}",
        f"Designation: {claimant_info.get('designation') or ''}    "
        f"Department: {claimant_info.get('department') or ''}    "
        f"Basic Pay (Rs.): {claimant_info.get('basic_pay') or ''}",
        f"Patient Name: {patient.get('name') or ''}    "
        f"S/o, W/o, D/o: {patient.get('father_or_husband_name') or ''}    "
        f"Relation: {patient.get('relation_to_claimant') or ''}",
        f"Hospital/Dispensary: {hospital.get('name') or ''}",
        f"Period of Treatment: {_fmt_date(admission.get('admission_date'))} "
        f"to {_fmt_date(admission.get('discharge_date'))}",
        f"Indoor No.: {admission.get('indoor_no') or ''}    "
        f"Admission Date: {_fmt_date(admission.get('admission_date'))}",
        f"He/She is suffering from: {admission.get('diagnosis') or ''}",
    ]
    cell = header_table.rows[0].cells[0]
    para = cell.paragraphs[0]
    fmt_run = para.runs[0] if para.runs else None
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for i, line in enumerate(header_lines):
        target_para = para if i == 0 else cell.add_paragraph()
        run = target_para.add_run(line)
        if fmt_run is not None:
            run.font.name = fmt_run.font.name
            run.font.size = fmt_run.font.size
            run.font.bold = fmt_run.font.bold

    # ---------- Item table ----------
    item_table = _find_item_table(doc)
    rows = _apply_ditto(_build_ec_rows(claim.get("items", [])))

    # find first data row (row after header) to use as formatting template,
    # then remove any pre-existing empty data rows before filling
    while len(item_table.rows) > 1:
        item_table._tbl.remove(item_table.rows[1]._tr)

    template_data_row = None
    if len(item_table.rows) == 1:
        # duplicate header row once to get a data-row template, then style as data row
        template_row = _clone_row(item_table, template_row_index=0)
        template_data_row = template_row
    section_labels_used = set(claim.get("items", []) and [i["category"] for i in claim["items"]])

    sr_no = 1
    for r in rows:
        new_row = _clone_row(item_table, template_row_index=1 if template_data_row else 0)
        cells = new_row.cells
        _set_cell_text(cells[0], sr_no)
        _set_cell_text(cells[1], r["name"])
        _set_cell_text(cells[2], r["bill_no_display"])
        _set_cell_text(cells[3], r["date_display"])
        _set_cell_text(cells[4], f'{r["amount"]:.2f}' if isinstance(r["amount"], (int, float)) else (r["amount"] or ""))
        sr_no += 1

    # remove the leftover formatting-template row (kept purely for run styling)
    if template_data_row is not None:
        item_table._tbl.remove(template_data_row._tr)

    # try to make the header row repeat on every printed page (Word feature)
    try:
        trPr = item_table.rows[0]._tr.get_or_add_trPr()
        tblHeader = trPr.makeelement(qn("w:tblHeader"), {})
        trPr.append(tblHeader)
    except Exception:
        pass

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
